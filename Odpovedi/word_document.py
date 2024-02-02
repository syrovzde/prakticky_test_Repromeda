from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import sys
import os

if __name__ == "__main__":
    """Only basic error handling"""
    if len(sys.argv) != 5:
        print("Usage: python script.py jmeno prijmeni rodne_cislo datum")
        sys.exit(1)
    doc = Document()
    text = "Výsledný protokol genetického vyšetření"
    title_paragraph = doc.add_paragraph(text)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.runs[0]
    title_run.font.size = Pt(16)  # Set font size to 16 points
    title_run.font.bold = True  # Set bold style
    title_run.font.color.rgb = RGBColor(0, 0, 0) 

    table = doc.add_table(rows=0, cols=2)
    data = [
        ["Jméno a příjmení:", sys.argv[1]+' ' + sys.argv[2]],
        ["Rodné číslo:", sys.argv[3]],
        ["Datum odběru:", sys.argv[4]]    
    ]   

    for (id,name) in data:
        row = table.add_row().cells
        row[0].text = id 
        row[1].text = name

    table.style = "Table Grid" 
    doc.save("generated.docx")